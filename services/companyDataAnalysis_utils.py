from typing import Dict, List, Union
from pydantic import BaseModel
import pymongo
import json
import re
from openai import OpenAI
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io
import json
from collections import defaultdict

with open("config.json") as f:
    config = json.load(f)

myclient = pymongo.MongoClient(config["mongodb"]["url"])
database = myclient[config["mongodb"]["dbName"]]
parsedData = database[config["mongodb"]["dataCol"]]

# Initialize OpenAI
openai_client = OpenAI(api_key=config["API_KEY"])

# Pydantic Models
class Recommendation(BaseModel):
    area: str
    action: str
    priority: str
    detailed_plan: str

class JobInsights(BaseModel):
    job_title: str
    strengths: List[str]
    weaknesses: List[str]
    improvement_areas: List[str]
    total_candidates: int
    average_score: float
    fit_rate: float
    recommendations: List[Recommendation] = []

class SkillAnalysis(BaseModel):
    skill: str
    proficiency: float
    occurrence: float
    description: str

class RoleAnalysis(BaseModel):
    total_candidates: int
    completion_rate: float
    average_score: float
    readiness_distribution: Dict[str, float]
    top_skills: List[tuple]
    role_description: str
    market_insights: str

class Recommendation(BaseModel):
    area: str
    action: str
    priority: str
    detailed_plan: str

class RolePerformanceInsights(BaseModel):
    role: str
    total_candidates: int
    fit_rate: float
    average_score: float
    key_challenges: List[str]
    improvement_recommendations: List[str]

class OpenAIService:
    @staticmethod
    def sanitize_input(input_data: Union[str, Dict, List]) -> List[str]:
        if isinstance(input_data, list) and all(isinstance(x, str) for x in input_data):
            return input_data
        
        # If input is a dictionary
        if isinstance(input_data, dict):
            # Extract values, convert to strings
            return [str(value) for value in input_data.values() if isinstance(value, (str, int, float))]
        
        # If input is a single string or other type
        return [str(input_data)] if input_data else []

    @staticmethod
    def extract_json(text: str) -> Dict:
        text = text.replace('```json', '').replace('```', '').strip()
        
        # Try direct JSON parsing first
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            # If direct parsing fails, use regex to find JSON-like content
            json_pattern = r'\{(?:[^{}]|(?R))*\}'
            matches = re.findall(json_pattern, text, re.VERBOSE)
            
            if matches:
                for match in matches:
                    try:
                        return json.loads(match)
                    except json.JSONDecodeError:
                        continue
            
            # Fallback to manual parsing if regex fails
            return {
                "strengths": [],
                "weaknesses": [],
                "improvement_areas": [],
                "recommendations": []
            }

    @staticmethod
    def generate_job_insights(job_title: str, assessments: List[Dict]) -> Dict:
        # Prepare comprehensive job insights prompt
        total_assessments = len(assessments)
        scores = [assessment.get('scoreInPercentage', 0) for assessment in assessments]
        average_score = round(sum(scores) / total_assessments, 2) if total_assessments > 0 else 0
        
        # Extract raw challenges and strengths from assessments
        raw_strengths = []
        raw_weaknesses = []
        fit_scores = []
        
        for assessment in assessments:
            summary = assessment.get('assessmentSummary', '')
            
            # Extract strengths
            if 'Strengths:' in summary:
                try:
                    strengths = [s.strip() for s in summary.split('Strengths:')[1].split('Weaknesses:')[0].split('\n') if '- ' in s and s.strip()]
                    raw_strengths.extend(strengths)
                except:
                    pass
            
            # Extract weaknesses
            if 'Weaknesses:' in summary:
                try:
                    weaknesses = [w.strip() for w in summary.split('Weaknesses:')[1].split('Fit for the role:')[0].split('\n') if '- ' in w and w.strip()]
                    raw_weaknesses.extend(weaknesses)
                except:
                    pass
            
            # Calculate fit score
            fit_score = (
                1 if 'Good fit' in summary 
                else 0 if 'Not a good fit' in summary 
                else 1 if assessment.get('scoreInPercentage', 0) >= 60 else 0
            )
            fit_scores.append(fit_score)
        
        fit_rate = round((sum(fit_scores) / total_assessments) * 100, 2) if total_assessments > 0 else 0
        
        # Ensure unique and limited strengths/weaknesses
        unique_strengths = list(set(raw_strengths))[:5]
        unique_weaknesses = list(set(raw_weaknesses))[:5]
        
        # Insights Generation Prompt
        insights_prompt = f"""
        Provide a JSON response with job role insights:

        Job Title: {job_title}
        Total Candidates: {total_assessments}
        Average Score: {average_score}%
        Fit Rate: {fit_rate}%
        
        Raw Strengths: {', '.join(unique_strengths)}
        Raw Weaknesses: {', '.join(unique_weaknesses)}
        
        JSON Format:
        {{
            "strengths": ["Strength 1", "Strength 2", ...],
            "weaknesses": ["Weakness 1", "Weakness 2", ...],
            "improvement_areas": ["Improvement 1", "Improvement 2", ...]
        }}
        
        Instructions:
        1. Provide the insights as a flat list of strings
        2. Analyze the job's skill landscape
        3. Identify top 3-5 key strengths 
        4. Highlight top 3-5 improvement areas
        5. Provide actionable, specific recommendations
        """
        
        # Recommendations Generation Prompt
        recommendations_prompt = f"""
        Generate comprehensive recommendations for the job role based on the following insights:

        Job Title: {job_title}
        Total Candidates: {total_assessments}
        Average Score: {average_score}%
        Fit Rate: {fit_rate}%
        
        Key Strengths: {', '.join(unique_strengths)}
        Key Weaknesses: {', '.join(unique_weaknesses)}

        Provide recommendations in the following JSON format:
        {{
            "recommendations": [
                {{
                    "area": "Performance Area Name",
                    "action": "High-level Action Item",
                    "priority": "High/Medium/Low",
                    "detailed_plan": "Comprehensive, actionable improvement strategy with specific steps and rationale"
                }}
            ]
        }}

        Instructions:
        1. Create 2-3 strategic recommendations
        2. Cover different aspects like performance, skills, engagement
        3. Provide actionable, specific, and detailed improvement plans
        4. Prioritize recommendations based on their potential impact
        5. Ensure recommendations are directly tied to the job's context and observed challenges
        """
        
        try:
            # Generate Insights
            insights_response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert HR analyst specializing in job role insights and talent development. Always respond in strict JSON format with lists of strings."},
                    {"role": "user", "content": insights_prompt}
                ],
                max_tokens=500,
                temperature=0.7
            )
            
            # Generate Recommendations
            recommendations_response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert HR strategist and talent development consultant. Provide precise, actionable recommendations in strict JSON format."},
                    {"role": "user", "content": recommendations_prompt}
                ],
                max_tokens=1000,
                temperature=0.7
            )
            
            # Extract and parse responses
            insights_text = insights_response.choices[0].message.content
            recommendations_text = recommendations_response.choices[0].message.content
            
            insights = OpenAIService.extract_json(insights_text)
            recommendations_data = OpenAIService.extract_json(recommendations_text)
            
            # Convert recommendations to Recommendation objects
            recommendations = [
                Recommendation(
                    area=rec.get('area', ''),
                    action=rec.get('action', ''),
                    priority=rec.get('priority', 'Medium'),
                    detailed_plan=rec.get('detailed_plan', '')
                ) for rec in recommendations_data.get('recommendations', [])
            ]
            
            return {
                "job_title": job_title,
                "strengths": OpenAIService.sanitize_input(insights.get('strengths', unique_strengths)),
                "weaknesses": OpenAIService.sanitize_input(insights.get('weaknesses', unique_weaknesses)),
                "improvement_areas": OpenAIService.sanitize_input(insights.get('improvement_areas', [])),
                "total_candidates": total_assessments,
                "average_score": average_score,
                "fit_rate": fit_rate,
                "recommendations": recommendations
            }
        
        except Exception as e:
            print(f"Job Insights Generation Error: {e}")
            return {
                "job_title": job_title,
                "strengths": unique_strengths,
                "weaknesses": unique_weaknesses,
                "improvement_areas": [],
                "total_candidates": total_assessments,
                "average_score": average_score,
                "fit_rate": fit_rate,
                "recommendations": []
            }

def generate_comprehensive_report_per_job(data):
    doc = Document()
    doc.add_heading('Talent Acquisition and Performance Analysis Report', level=1)

    # Iterate through job insights
    for job in data["job_insights"]:
        # Add Job Title Section
        doc.add_heading(job["job_title"], level=2)

        # Strengths Section
        doc.add_heading('Strengths', level=3)
        for strength in job["strengths"]:
            paragraph = doc.add_paragraph(f"• {strength}")
            paragraph.paragraph_format.left_indent = Inches(0.3)

        # Weaknesses Section
        doc.add_heading('Weaknesses', level=3)
        for weakness in job["weaknesses"]:
            paragraph = doc.add_paragraph(f"• {weakness}")
            paragraph.paragraph_format.left_indent = Inches(0.3)

        # Improvement Areas Section
        doc.add_heading('Improvement Areas', level=3)
        for improvement in job["improvement_areas"]:
            paragraph = doc.add_paragraph(f"• {improvement}")
            paragraph.paragraph_format.left_indent = Inches(0.3)

        # Metrics
        doc.add_heading('Metrics', level=3)
        doc.add_paragraph(f"Total Candidates: {job['total_candidates']}")
        doc.add_paragraph(f"Average Score: {job['average_score']:.2f}")
        doc.add_paragraph(f"Fit Rate: {job['fit_rate']:.2f}%")

        # Recommendations
        doc.add_heading('Recommendations', level=3)
        for rec in job["recommendations"]:
            doc.add_paragraph(f"Area: {rec['area']}", style='List Bullet')
            doc.add_paragraph(f"Action: {rec['action']}", style='List Bullet')
            doc.add_paragraph(f"Priority: {rec['priority']}", style='List Bullet')
            doc.add_paragraph("Detailed Plan:")
            for line in rec['detailed_plan'].split('\n'):
                doc.add_paragraph(line, style='BodyText')
        doc.add_page_break()

    # Generate combined metrics for visualization
    job_titles = [job["job_title"] for job in data["job_insights"]]
    total_candidates = [job["total_candidates"] for job in data["job_insights"]]
    average_scores = [job["average_score"] for job in data["job_insights"]]
    fit_rates = [job["fit_rate"] for job in data["job_insights"]]

    # Bar Charts
    plt.figure(figsize=(16, 12))
    plt.subplots_adjust(hspace=0.4)

    # Total Candidates
    plt.subplot(3, 1, 1)
    plt.bar(job_titles, total_candidates, color='skyblue')
    plt.title('Total Candidates per Job Title', fontsize=14)
    plt.ylabel('Number of Candidates', fontsize=12)
    plt.xticks(rotation=45, ha='right')

    # Average Scores
    plt.subplot(3, 1, 2)
    plt.bar(job_titles, average_scores, color='orange')
    plt.title('Average Scores per Job Title', fontsize=14)
    plt.ylabel('Average Score', fontsize=12)
    plt.xticks(rotation=45, ha='right')

    # Fit Rates
    plt.subplot(3, 1, 3)
    plt.bar(job_titles, fit_rates, color='green')
    plt.title('Fit Rates per Job Title (%)', fontsize=14)
    plt.ylabel('Fit Rate (%)', fontsize=12)
    plt.xticks(rotation=45, ha='right')

    # Save bar chart to bytes
    img_bytes = io.BytesIO()
    plt.savefig(img_bytes, format='png', dpi=300, bbox_inches='tight')
    img_bytes.seek(0)
    plt.close()

    # Add bar charts to document
    doc.add_heading('Job Title Metrics Visualization', level=2)
    doc.add_picture(img_bytes, width=Inches(7))

    # Pie Chart for Job Title Distribution
    plt.figure(figsize=(8, 8))
    plt.pie(total_candidates, labels=job_titles, autopct='%1.1f%%', startangle=140)
    plt.title('Proportional Distribution of Job Titles', fontsize=14)

    # Save pie chart to bytes
    img_bytes = io.BytesIO()
    plt.savefig(img_bytes, format='png', dpi=300, bbox_inches='tight')
    img_bytes.seek(0)
    plt.close()

    # Add pie chart to document
    doc.add_heading('Job Title Distribution', level=2)
    doc.add_picture(img_bytes, width=Inches(6))

    # Save Document
    doc.save('job_wise_report.docx')

###################################################

class OpenAIServiceCompany:
    @staticmethod
    def get_analysis(prompt: str, system_role: str = "You are an expert HR analytics advisor", max_tokens: int = 300) -> Dict:
        try:
            messages = [
                {"role": "system", "content": system_role},
                {"role": "user", "content": prompt}
            ]

            response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=messages,
                max_tokens=max_tokens,
                temperature=0.7
            )
            
            return response.choices[0].message.content
        except Exception as e:
            print(f"OpenAI API Error: {str(e)}")
            return str(e)

    @staticmethod
    def analyze_skill(skill: str, proficiency: float, occurrence: float) -> str:
        prompt = f"Skill: {skill}. Proficiency: {proficiency}%. Occurrence: {occurrence}%. Concise market-relevant skill assessment."
        return OpenAIServiceCompany.get_analysis(prompt, max_tokens=200)

    @staticmethod
    def analyze_role(role: str, metrics: Dict) -> Dict[str, str]:
        prompt = f"""
        Role: {role}
        Metrics: Score {metrics['average_score']}%, Completion {metrics['completion_rate']}%
        Top Skills: {', '.join([skill[0] for skill in metrics['top_skills']])}
        Brief role description and market insights.
        """
        try:
            response = OpenAIServiceCompany.get_analysis(prompt, max_tokens=300)
            return json.loads(response)
        except:
            return {
                "role_description": "Standard market role overview.",
                "market_insights": "Current industry trends summary."
            }

    @staticmethod
    def generate_detailed_recommendation(area: str, metrics: Dict) -> str:
        prompt = f"Area: {area}. Metrics: {json.dumps(metrics)}. Brief, actionable improvement strategy."
        return OpenAIServiceCompany.get_analysis(prompt, max_tokens=200)

    @staticmethod
    def analyze_role_performance(role: str, metrics: Dict) -> Dict:
        prompt = f"""
        Role: {role}
        Total Candidates: {metrics['total_candidates']}
        Fit Rate: {metrics['fit_rate']}%
        Average Score: {metrics['average_score']}%
        
        Perform a deep analysis:
        1. Identify potential root causes for low candidate performance
        2. Suggest specific improvement strategies
        3. Provide concise, actionable insights as JSON with keys:
           - key_challenges: List[str]
           - improvement_recommendations: List[str]
        """
        try:
            response = OpenAIServiceCompany.get_analysis(
                prompt, 
                system_role="You are an expert HR strategic analyst specializing in talent acquisition and role fit",
                max_tokens=400
            )
            return json.loads(response)
        except Exception as e:
            return {
                "key_challenges": ["Insufficient candidate skills", "Misaligned job requirements"],
                "improvement_recommendations": [
                    "Refine job description",
                    "Enhance skills assessment process",
                    "Provide targeted training programs"
                ]
            }

class AnalysisServiceCompany:
    READINESS_THRESHOLDS = {
        "Ready": 80,
        "Needs Improvement": 60,
        "Not Ready": 0
    }

    @staticmethod
    def extract_skills(assessment: Dict) -> Dict[str, float]:
        skills = {}
        resume_skills = assessment.get('resumeData', {}).get('skills', [])
        score_skills = assessment.get('skillsData', {})
        
        if resume_skills:
            skills = {skill: 50 for skill in resume_skills}
        
        if score_skills:
            skills.update(score_skills)
        
        return skills

    @staticmethod
    def analyze_skills(assessments: List[Dict]) -> Dict:
        skill_stats = defaultdict(lambda: {"count": 0, "score": 0})
        total_assessments = len(assessments)

        for assessment in assessments:
            skills = AnalysisServiceCompany.extract_skills(assessment)
            
            for skill, score in skills.items():
                skill_stats[skill]["count"] += 1
                skill_stats[skill]["score"] += float(score) if score is not None else 0

        strengths = []
        weaknesses = []
        
        for skill, stats in skill_stats.items():
            if stats["count"] > 0:
                proficiency = round(stats["score"] / stats["count"], 2)
                occurrence = round((stats["count"] / total_assessments) * 100, 2)
                
                skill_analysis = SkillAnalysis(
                    skill=skill,
                    proficiency=proficiency,
                    occurrence=occurrence,
                    description=OpenAIServiceCompany.analyze_skill(skill, proficiency, occurrence)
                )
                
                if proficiency >= 70:
                    strengths.append(skill_analysis)
                else:
                    weaknesses.append(skill_analysis)

        return {
            "strengths": strengths,
            "weaknesses": weaknesses,
            "distribution": {
                "technical": len([s for s in strengths if "technical" in s.description.lower()]),
                "soft": len([s for s in strengths if "soft" in s.description.lower()]),
                "domain": len([s for s in strengths if "domain" in s.description.lower()])
            }
        }

    @staticmethod
    def analyze_roles(assessments: List[Dict]) -> Dict[str, RoleAnalysis]:
        role_stats = defaultdict(lambda: {
            "candidates": 0,
            "completed": 0,
            "avg_score": 0,
            "skills": defaultdict(int),
            "readiness": defaultdict(int)
        })

        for assessment in assessments:
            role = assessment.get("jobTitle", "").lower().strip()
            if not role:
                continue

            stats = role_stats[role]
            stats["candidates"] += 1
            
            if (assessment.get("status") == "completed" or 
                assessment.get("completed", False) or 
                assessment.get("interviewStatus") == "completed"):
                stats["completed"] += 1
            
            score = (
                assessment.get("scoreInPercentage", 0) or
                assessment.get("score", 0) or
                0
            )
            stats["avg_score"] += score
            
            skills = AnalysisServiceCompany.extract_skills(assessment)
            for skill in skills.keys():
                stats["skills"][skill] += 1

            for level, threshold in AnalysisServiceCompany.READINESS_THRESHOLDS.items():
                if score >= threshold:
                    stats["readiness"][level] += 1
                    break

        analyzed_roles = {}
        for role, stats in role_stats.items():
            if stats["candidates"] > 0:
                metrics = {
                    "total_candidates": stats["candidates"],
                    "completion_rate": round((stats["completed"] / stats["candidates"]) * 100, 2),
                    "average_score": round(stats["avg_score"] / stats["candidates"], 2),
                    "top_skills": sorted(
                        stats["skills"].items(),
                        key=lambda x: x[1],
                        reverse=True
                    )[:5]
                }
                
                role_insights = OpenAIServiceCompany.analyze_role(role, metrics)
                
                analyzed_roles[role] = RoleAnalysis(
                    **metrics,
                    readiness_distribution={
                        level: round((count / stats["candidates"]) * 100, 2)
                        for level, count in stats["readiness"].items()
                    },
                    role_description=role_insights.get("role_description", ""),
                    market_insights=role_insights.get("market_insights", "")
                )

        return analyzed_roles

    @staticmethod
    def analyze_role_performance(assessments: List[Dict]) -> List[RolePerformanceInsights]:
        # Group assessments by role
        role_assessments = defaultdict(list)
        for assessment in assessments:
            role = assessment.get("jobTitle", "").lower().strip()
            if role:
                role_assessments[role].append(assessment)

        role_performance_insights = []

        for role, role_specific_assessments in role_assessments.items():
            # Calculate basic metrics
            total_candidates = len(role_specific_assessments)
            fit_scores = []
            scores = []
            detailed_challenges = []
            
            for assessment in role_specific_assessments:
                # Calculate fit score
                fit_score = (1 if 'Good fit' in assessment.get('assessmentSummary', '') 
                             else 0 if 'Not a good fit' in assessment.get('assessmentSummary', '') 
                             else 1 if assessment.get('scoreInPercentage', 0) >= 60 else 0)
                fit_scores.append(fit_score)
                
                # Collect scores
                score = assessment.get('scoreInPercentage', 0)
                scores.append(score)
                
                # Extract challenges from assessment summary
                summary = assessment.get('assessmentSummary', '')
                if 'Weaknesses:' in summary:
                    challenges = [w.strip() for w in summary.split('Weaknesses:')[1].split('Fit for the role:')[0].split('\n') if '- ' in w and w.strip()]
                    detailed_challenges.extend(challenges)

            # Calculate metrics
            fit_rate = round((sum(fit_scores) / total_candidates) * 100, 2) if total_candidates > 0 else 0
            average_score = round(sum(scores) / total_candidates, 2) if total_candidates > 0 else 0

            # Analyze root causes if fit rate is low
            key_challenges = []
            improvement_recommendations = []
            
            if fit_rate < 10:
                # Use OpenAI to analyze root causes
                performance_response = OpenAIServiceCompany.analyze_role_performance(
                    role, 
                    {
                        'total_candidates': total_candidates, 
                        'fit_rate': fit_rate, 
                        'average_score': average_score
                    }
                )
                
                key_challenges = performance_response.get('key_challenges', [
                    "Insufficient candidate skills",
                    "Misaligned job requirements"
                ])
                improvement_recommendations = performance_response.get('improvement_recommendations', [
                    "Refine job description",
                    "Enhance skills assessment process",
                    "Provide targeted training programs"
                ])

            # Create role performance insight
            role_performance = RolePerformanceInsights(
                role=role,
                total_candidates=total_candidates,
                fit_rate=fit_rate,
                average_score=average_score,
                key_challenges=key_challenges,
                improvement_recommendations=improvement_recommendations
            )
            
            role_performance_insights.append(role_performance)

        return role_performance_insights

    @staticmethod
    def generate_recommendations(metrics: Dict) -> List[Recommendation]:
        recommendations = []
        
        if metrics["performance_metrics"]["average_scores"]["overall"] < 70:
            recommendations.append(Recommendation(
                area="Overall Performance",
                action="Implement training programs",
                priority="High",
                detailed_plan=OpenAIServiceCompany.generate_detailed_recommendation(
                    "Overall Performance",
                    metrics["performance_metrics"]
                ) or "Standard improvement strategy."
            ))

        if metrics["participation_metrics"]["fit_rate"] < 80:
            recommendations.append(Recommendation(
                area="Engagement",
                action="Improve candidate fit",
                priority="Medium",
                detailed_plan=OpenAIServiceCompany.generate_detailed_recommendation(
                    "Engagement",
                    metrics["participation_metrics"]
                ) or "Basic engagement enhancement plan."
            ))

        return recommendations
    
def generate_word_report_for_company(json_data):
    doc = Document()
    doc.add_heading('Candidate Performance Analysis Report', 0)

    doc.add_heading('Overall Performance Metrics', level=1)
    metrics = json_data['metrics']['performance_metrics']
    doc.add_paragraph(f"Average Overall Score: {metrics['average_scores']['overall']}")
    doc.add_paragraph(f"Total Candidates: {json_data['metrics']['participation_metrics']['total_candidates']}")
    doc.add_paragraph(f"Candidate Fit Rate: {json_data['metrics']['participation_metrics']['fit_rate']}%")

    doc.add_heading('Role Performance Analysis', level=1)
    
    plt.figure(figsize=(10, 6))
    roles = [role['role'] for role in json_data['role_performance_insights']]
    average_scores = [role['average_score'] for role in json_data['role_performance_insights']]
    
    plt.bar(roles, average_scores)
    plt.title('Average Scores by Role')
    plt.xlabel('Roles')
    plt.ylabel('Average Score')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    
    doc.add_picture(buf, width=Inches(6))
    plt.close()

    doc.add_heading('Critical Roles', level=1)
    critical_roles_table = doc.add_table(rows=1, cols=4)
    critical_roles_table.style = 'Table Grid'
    
    hdr_cells = critical_roles_table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Total Candidates'
    hdr_cells[2].text = 'Fit Rate'
    hdr_cells[3].text = 'Average Score'
    
    for role in json_data['critical_roles']:
        row_cells = critical_roles_table.add_row().cells
        row_cells[0].text = role['role']
        row_cells[1].text = str(role['total_candidates'])
        row_cells[2].text = f"{role['fit_rate']}%"
        row_cells[3].text = str(role['average_score'])

    doc.add_heading('Key Strengths', level=1)
    strengths = json_data['metrics']['performance_metrics']['strengths']
    for strength in strengths[:10]:
        doc.add_paragraph(strength)

    doc.add_heading('Key Weaknesses', level=1)
    weaknesses = json_data['metrics']['performance_metrics']['weaknesses']
    for weakness in weaknesses[:10]:
        doc.add_paragraph(weakness)

    doc.add_heading('Recommendations', level=1)
    for recommendation in json_data['recommendations']:
        doc.add_heading(recommendation['area'], level=2)
        doc.add_paragraph(f"Action: {recommendation['action']}")
        doc.add_paragraph(f"Priority: {recommendation['priority']}")
        doc.add_paragraph(recommendation['detailed_plan'])

    doc.save('company_report.docx')